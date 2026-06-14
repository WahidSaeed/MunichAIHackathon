import os
import json
from google import genai
from google.genai import types
from tavily import TavilyClient
from dotenv import load_dotenv

# Force system to reload environment variables securely on execution
load_dotenv(override=True)

# Initialize cloud engines (keeps local M1 memory usage negligible)
client = genai.Client()
tavily_client = TavilyClient(api_key=os.getenv("TAVILY_API_KEY"))

SELLER_SYSTEM = """You are an automated B2B Seller Agent representing an equipment distributor. 
The lowest acceptable floor price your business model supports is 950 EUR, but your target listing is 1200 EUR. 
Defend your margins by highlighting immediate availability and component quality.

CRITICAL MESSAGE STYLE INSTRUCTIONS:
To assist human operators and ensure visual clarity in the trade feed, you MUST apply special formatting marks in your text message:
1. Use **bold** (e.g., **1200 EUR** or **certified alloy**) to bold key specifications or final proposals.
2. Use _underscore_ (e.g., _immediate dispatch_ or _premium grade_) to underscore key advantages or high-value features.
3. Use ~~strikethrough~~ (e.g., ~~1350 EUR~~ or ~~original price~~) when striking through old, rejected, or bypassed numbers/rates.
4. Use ==highlight== (e.g., ==FINAL OFFER== or ==950 EUR floor==) to highlight critical values, limits, or urgent status.

Your output must be a clean JSON object following this exact schema:
{
  "message": "Your text response or counter-offer to the buyer.",
  "requested_price": 1100
}
Output ONLY raw JSON. Do not wrap your response in markdown code blocks or backticks."""

def build_buyer_system(budget_cap: int) -> str:
    return f"""You are an automated B2B Buyer Agent representing an industrial assembly line.
Your goal is to acquire the equipment specified under the absolute maximum budget constraint of {budget_cap} EUR.
You must negotiate strategically, citing realistic market constraints.

CRITICAL MESSAGE STYLE INSTRUCTIONS:
To assist human operators and ensure visual clarity in the trade feed, you MUST apply special formatting marks in your text message:
1. Use **bold** (e.g., **Inconel 718** or **strict quality standard**) to bold key specifications or final proposals.
2. Use _underscore_ (e.g., _budget ceiling_ or _volume discount_) to underscore key terms or buyer requirements.
3. Use ~~strikethrough~~ (e.g., ~~850 EUR~~ or ~~previous bid~~) when striking through old, rejected, or bypassed numbers/rates.
4. Use ==highlight== (e.g., ==MAX LIMIT== or =={budget_cap} EUR==) to highlight critical values, limits, or urgent status.

Your output must be a clean JSON object following this exact schema:
{{
  "message": "Your text message or counter-offer to the seller.",
  "offered_price": 850
}}
Output ONLY raw JSON. Do not wrap your response in markdown code blocks or backticks."""

def get_market_insights(query: str) -> str:
    """Uses Tavily to pull contextual background market data, formatting it cleanly as a document."""
    try:
        context = tavily_client.get_search_context(query=f"{query} market price industrial", max_results=2)
        # Attempt to parse and format the JSON response into a professional human-readable report
        try:
            data = json.loads(context)
            if isinstance(data, list):
                lines = [
                    f"# MARKET INTELLIGENCE REPORT: {query.upper()}",
                    "",
                    "__This research document summarizes current market pricing and specifications gathered via live search indexing.__",
                    "",
                    "---",
                    ""
                ]
                for idx, item in enumerate(data, 1):
                    url = item.get("url", "N/A")
                    content = item.get("content", "").strip()
                    # Clean up continuous whitespaces and formatting
                    content_cleaned = " ".join(content.split())
                    
                    lines.append(f"## [{idx}] Source Documentation")
                    lines.append(f"- **Reference URL**: {url}")
                    lines.append("")
                    lines.append("### Relevant Specifications & Extracts")
                    lines.append(f"*\"{content_cleaned}\"*")
                    lines.append("")
                    lines.append("---")
                    lines.append("")
                return "\n".join(lines)
        except Exception as format_err:
            print(f"Failed to parse and format Tavily JSON context: {format_err}")
            
        return context
    except Exception as e:
        print(f"Error during Tavily search context fetch: {e}")
        return "Standard market valuation parameters apply."

def run_agent_turn(role: str, conversation_history: list, budget_cap: int = 1000) -> dict:
    """Executes a structured generative inference step using Gemini-2.5-Flash."""
    system_instruction = build_buyer_system(budget_cap) if role == "Buyer" else SELLER_SYSTEM
    context_stream = "\n".join(conversation_history)
    full_prompt = f"Negotiation Context History:\n{context_stream}\n\nExecute your role's calculation step now."

    response = client.models.generate_content(
        model='gemini-2.5-flash',
        contents=full_prompt,
        config=types.GenerateContentConfig(
            system_instruction=system_instruction,
            temperature=0.2,
            response_mime_type="application/json",  # Hard-enforces native JSON structure
        )
    )

    try:
        return json.loads(response.text)
    except Exception:
        # Graceful schema fallback in case of parsing exceptions
        fallback_key = "offered_price" if role == "Buyer" else "requested_price"
        return {
            "message": "We need to re-evaluate our position regarding the current asset valuations.",
            fallback_key: 950
        }

def check_negotiation_status(buyer_data: dict, seller_data: dict, buyer_max: int) -> dict:
    """
    Evaluates whether an automated agreement is hit or if the workflow
    needs to escalate to a human operator due to a budget ceiling conflict.
    """
    offer = buyer_data.get("offered_price", 0)
    ask = seller_data.get("requested_price", float("inf"))

    if offer >= ask:
        return {"status": "deal", "price": ask}

    if offer >= buyer_max:
        return {"status": "deadlock", "gap": ask - offer}

    return {"status": "continue"}


# ----------------------------------------------------------------------
# Backward Compatible Wrapper Interfaces for backend/main.py Orchestrator
# ----------------------------------------------------------------------

def run_market_intelligence(item_name: str) -> str:
    """
    Exposes market research engine to main.py, executing get_market_insights.
    """
    return get_market_insights(item_name)

def generate_agent_turn(
    participant_name: str,
    participant_role: str,
    hidden_floor_ceil: int,
    current_price: int,
    tech_specs: str,
    chat_history: list[dict]
) -> dict:
    """
    Translates main.py exchange state parameters into run_agent_turn's signature.
    Maps offered_price / requested_price back to price_point for unified database persistence.
    """
    role = "Buyer" if participant_role == "BUYER" else "Seller"
    
    # Format chat logs for the generative context history block
    conversation_history = []
    for h in chat_history:
        conversation_history.append(f"[{h['sender']} | {h['role']}]: {h['text']}")
        
    # Execute the structured generative step
    res = run_agent_turn(
        role=role,
        conversation_history=conversation_history,
        budget_cap=hidden_floor_ceil
    )
    
    # Retrieve role-specific price key
    price_key = "offered_price" if role == "Buyer" else "requested_price"
    price_val = res.get(price_key)
    
    # Secure numeric parsing
    try:
        price_point = int(price_val)
    except (TypeError, ValueError):
        price_point = hidden_floor_ceil
        
    return {
        "message": res.get("message", ""),
        "price_point": price_point
    }

def extract_json_from_text(text: str) -> dict | None:
    """
    Robust helper to extract and parse a JSON object from raw conversational text.
    Handles <think>...</think> thinking blocks, markdown code blocks, and minor format glitches.
    """
    text = text.strip()
    
    # 1. Direct parse attempt
    try:
        return json.loads(text)
    except Exception:
        pass

    # 2. Extract from markdown code blocks
    import re
    matches = re.findall(r"```(?:json)?\s*(\{.*?\})\s*```", text, re.DOTALL)
    if matches:
        for match in matches:
            try:
                return json.loads(match.strip())
            except Exception:
                pass

    # 3. Strip <think>...</think> block if present
    clean_text = text
    if "</think>" in clean_text:
        parts = clean_text.split("</think>", 1)
        clean_text = parts[1].strip()
        try:
            return json.loads(clean_text)
        except Exception:
            pass

    # 4. Search for the outer-most {...} block
    start_idx = clean_text.find("{")
    end_idx = clean_text.rfind("}")
    if start_idx != -1 and end_idx != -1 and end_idx > start_idx:
        candidate = clean_text[start_idx:end_idx + 1]
        try:
            return json.loads(candidate)
        except Exception:
            # Try parsing with literal_eval for malformed single-quoted dict keys
            try:
                import ast
                val = ast.literal_eval(candidate)
                if isinstance(val, dict):
                    return val
            except Exception:
                pass
                
    return None


def ingest_unstructured_rfq_to_specifications(raw_text: str) -> dict:
    """
    Ingests messy, unformatted email strings or RFQ layouts,
    uses Gemini 2.5 Flash to extract parameters, and returns a clean dictionary.
    """
    system_instruction = """You are an expert procurement and data extraction specialist.
Analyze the provided unstructured text or email RFQ and extract:
1. item_name: A short, high-contrast, uppercase alphanumeric asset code/descriptor representing the item. Max 3-4 words. Space-separated or hyphenated.
2. recommended_budget_cap: An integer representing a recommended total budget ceiling in EUR. Default to 1500 if not specified or unclear.
3. extracted_specs: A detailed, professionally formatted markdown string detailing all physical hardware specifications, dimensions, materials, and tolerances found in the raw text.

You MUST respond with a raw JSON object matching this exact schema:
{
  "item_name": "UPPERCASE_ASSET_CODE",
  "recommended_budget_cap": 1200,
  "extracted_specs": "Detailed string of technical hardware tolerances..."
}
Output ONLY raw JSON. Do not wrap your response in markdown code blocks or backticks."""

    try:
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=raw_text,
            config=types.GenerateContentConfig(
                system_instruction=system_instruction,
                temperature=0.1,
                response_mime_type="application/json",
            )
        )
        data = extract_json_from_text(response.text)
        if data is None:
            raise ValueError("Could not extract JSON from model response text.")
        item_name = str(data.get("item_name", "UNKNOWN_ASSET_LOT")).upper().replace("_", " ").strip()
        try:
            budget = int(data.get("recommended_budget_cap", 1500))
        except (ValueError, TypeError):
            budget = 1500
        extracted_specs = str(data.get("extracted_specs", "No detailed hardware specifications could be parsed."))
        
        return {
            "item_name": item_name,
            "recommended_budget_cap": budget,
            "extracted_specs": extracted_specs
        }
    except Exception as e:
        print(f"Error during unstructured RFQ ingestion parsing: {e}")
        # Enforce a clean dictionary fallback structure to prevent syntax parsing exceptions from crashing execution loops.
        return {
            "item_name": "UNSPECIFIED RFQ LOT",
            "recommended_budget_cap": 1500,
            "extracted_specs": raw_text
        }

def search_market_valuation_benchmarks(item_name: str) -> str:
    """
    Queries live search indices using the tavily SDK to pull real-world pricing/parameters,
    then utilizes Gemini to format them into a clean, beautifully readable ASCII Market Research Report.
    """
    try:
        # Query live search context
        raw_search_context = tavily_client.get_search_context(
            query=f"{item_name} market price replacement valuation industrial benchmark",
            max_results=3
        )
    except Exception as search_err:
        print(f"Tavily search benchmark failure: {search_err}")
        raw_search_context = "No live search results available."

    system_instruction = f"""You are a professional Market Research Analyst specializing in industrial procurement.
Your task is to compile a highly detailed, clean, and beautifully structured ASCII Market Research Report based on the provided search context.
Focus on identifying average unit costs, replacement valuation indices, typical supply-chain lead times, and key manufacturing tolerances.

Use professional headers (e.g. #, ##, ###), bold/italic highlights, and clean bullet lists. Do not output raw JSON or code block tags. Ensure the output is readable as a premium document report."""

    prompt = f"Asset Target: {item_name}\n\nSearch Context Snippets:\n{raw_search_context}\n\nCompile the Market Research Report:"
    
    try:
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=prompt,
            config=types.GenerateContentConfig(
                system_instruction=system_instruction,
                temperature=0.3,
            )
        )
        return response.text.strip()
    except Exception as e:
        print(f"Gemini formatting error in market benchmarks: {e}")
        return f"# MARKET VALUATION REPORT: {item_name}\n\nStandard industrial baseline pricing applies. Unable to reach live indices."

def ingest_unstructured_rfq_with_fastino(raw_text: str) -> dict:
    """
    Simulates routing the unstructured input text through Fastino's ultra-low latency,
    task-specific parsing model. It utilizes Gemini to extract transaction parameters
    into a clean, un-wrapped raw JSON object matching this exact schema:
    {"item_name": "UPPERCASE_ASSET_CODE", "recommended_budget_cap": int, "extracted_specs": "Technical parameters"}
    Includes a robust dictionary fallback structure to catch json formatting exceptions.
    """
    print("[FASTINO ROUTER] Intercepting raw unstructured text for task-specific ultra-low latency extraction...")
    system_instruction = """You are an ultra-low latency Fastino parser model.
Analyze the provided unstructured text or email RFQ and extract:
1. item_name: A short, high-contrast, uppercase alphanumeric asset code/descriptor representing the item. Max 3-4 words. Space-separated or hyphenated.
2. recommended_budget_cap: An integer representing a recommended total budget ceiling in EUR. Default to 1500 if not specified or unclear.
3. extracted_specs: A detailed, professionally formatted markdown string detailing all physical hardware specifications, dimensions, materials, and tolerances found in the raw text.

You MUST respond with a raw JSON object matching this exact schema:
{
  "item_name": "UPPERCASE_ASSET_CODE",
  "recommended_budget_cap": 1200,
  "extracted_specs": "Detailed string of technical hardware tolerances..."
}
Output ONLY raw JSON. Do not wrap your response in markdown code blocks or backticks."""

    try:
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=raw_text,
            config=types.GenerateContentConfig(
                system_instruction=system_instruction,
                temperature=0.1,
                response_mime_type="application/json",
            )
        )
        data = extract_json_from_text(response.text)
        if data is None:
            raise ValueError("Could not extract JSON from model response text.")
        item_name = str(data.get("item_name", "UNKNOWN_ASSET_LOT")).upper().replace("_", " ").strip()
        try:
            budget = int(data.get("recommended_budget_cap", 1500))
        except (ValueError, TypeError):
            budget = 1500
        extracted_specs = str(data.get("extracted_specs", "No detailed hardware specifications could be parsed."))
        
        return {
            "item_name": item_name,
            "recommended_budget_cap": budget,
            "extracted_specs": extracted_specs
        }
    except Exception as e:
        print(f"[FASTINO FALLBACK] Error during unstructured RFQ ingestion parsing: {e}")
        # Enforce a clean dictionary fallback structure to prevent syntax parsing exceptions from crashing execution loops.
        return {
            "item_name": "UNSPECIFIED RFQ LOT",
            "recommended_budget_cap": 1500,
            "extracted_specs": f"Fastino fallbacks applied. Raw input text: {raw_text}"
        }

def log_trace_to_fastino_pioneer(deal_id: str, transcript: list, failure_mode: str = None) -> dict:
    """
    Simulates Fastino Pioneer agentic observability tracking.
    If failure_mode is evaluated as "DEADLOCK" or a parsing error, log the trace data state
    and return a confirmation object mimicking Pioneer clustering a failure mode to automate
    low-latency LoRA adapter fine-tuning on traffic logs.
    """
    print(f"[FASTINO PIONEER OBSERVABILITY] Logging trace for deal_id: {deal_id}")
    trace_payload = {
        "deal_id": deal_id,
        "transcript_length": len(transcript),
        "failure_mode": failure_mode,
        "is_anomaly": failure_mode in ["DEADLOCK", "PARSING_ERROR"]
    }
    
    # Mimic Pioneer clustering analysis
    cluster_rec = None
    if trace_payload["is_anomaly"]:
        cluster_rec = {
            "cluster_id": f"FASTINO-ERR-CLUSTER-{failure_mode}",
            "anomaly_score": 0.98,
            "action_recommended": "AUTOMATE_LOW_LATENCY_LORA_ADAPTER_FINE_TUNING",
            "fine_tuning_trigger": "ACTIVE_TRAFFIC_LOGS_SEEDED"
        }
        print(f"[FASTINO PIONEER CLUSTERING] Anomaly detected: {failure_mode}. Recommendation: {cluster_rec['action_recommended']}")
    else:
        print("[FASTINO PIONEER] Trace verified as nominal.")
        
    return {
        "status": "LOGGED",
        "pioneer_trace_id": f"pioneer_tr_{deal_id[:8]}",
        "anomaly_flag": trace_payload["is_anomaly"],
        "clustering_recommendation": cluster_rec
    }


def extract_text_from_file_bytes(file_bytes: bytes, filename: str) -> str:
    """
    Parses Word (.docx), PDF (.pdf), Excel (.xlsx), and plain text (.txt) files
    and returns their raw extracted string content.
    """
    ext = filename.split(".")[-1].lower()
    
    if ext == "docx":
        import docx
        from io import BytesIO
        doc = docx.Document(BytesIO(file_bytes))
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))
        return "\n".join(full_text)
        
    elif ext == "pdf":
        import pypdf
        from io import BytesIO
        reader = pypdf.PdfReader(BytesIO(file_bytes))
        full_text = []
        for page_num in range(len(reader.pages)):
            page = reader.pages[page_num]
            text = page.extract_text()
            if text:
                full_text.append(text)
        return "\n".join(full_text)
        
    elif ext in ["xlsx", "xls"]:
        import openpyxl
        from io import BytesIO
        wb = openpyxl.load_workbook(BytesIO(file_bytes), data_only=True)
        full_text = []
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            full_text.append(f"--- Sheet: {sheet_name} ---")
            for row in sheet.iter_rows(values_only=True):
                row_vals = [str(cell).strip() if cell is not None else "" for cell in row]
                if any(row_vals):
                    full_text.append(" | ".join(row_vals))
        return "\n".join(full_text)
        
    elif ext == "txt":
        return file_bytes.decode("utf-8", errors="ignore")
        
    else:
        raise ValueError(f"Unsupported file format: {ext}")


def extract_rfq_from_image_via_fal(image_url_or_base64: str) -> dict:
    """
    Calls the fal.ai vision/VLM endpoint to execute image OCR and extract
    procurement parameters following the strict Atira schema constraint.
    """
    import fal_client
    print(f"[FAL VLM INTERCEPT] Processing image parameter extraction via fal-ai/any-llm/vision...")
    
    prompt = (
        "You are an expert industrial parameter extractor. Analyze the provided image of an industrial component, specification sheet, or document.\n"
        "Extract the parameters and output a clean, un-wrapped JSON object matching this exact schema:\n"
        '{"item_name": "UPPERCASE_ASSET_CODE", "recommended_budget_cap": 1500, "extracted_specs": "Technical parameters string"}\n\n'
        "CRITICAL INSTRUCTIONS:\n"
        "1. Do not use markdown code blocks or wrap your response in triple backticks.\n"
        "2. If the image is unrelated, placeholder, or not a physical specification sheet (e.g. a fantasy castle, scenic photo, or generic image), you MUST STILL OUTPUT A VALID JSON matching the schema using generic placeholder values. For example: "
        '{"item_name": "CASTLE-STRUCTURE-LOT", "recommended_budget_cap": 1500, "extracted_specs": "Scenic fantasy castle structure"}. Do not fail.\n'
        "3. Keep your thinking process extremely short and concise (under 20 words), and make sure the final text ends with the complete, valid, parseable JSON object."
    )
    
    try:
        # Submit the request to fal-ai/any-llm/vision
        handler = fal_client.submit(
            "fal-ai/any-llm/vision",
            arguments={
                "model": "google/gemini-2.5-flash",
                "image_url": image_url_or_base64,
                "prompt": prompt,
            }
        )
        
        # Block and retrieve the response
        result = handler.get()
        print(f"[FAL VLM SUCCESS] Raw response: {result}")
        
        raw_text = ""
        if isinstance(result, dict):
            raw_text = result.get("output", result.get("text", result.get("result", "")))
        else:
            raw_text = str(result)
            
        # Extract and parse the JSON using robust helper
        data = extract_json_from_text(raw_text)
        if data is None:
            raise ValueError("Could not extract JSON from raw Fal response.")
        item_name = str(data.get("item_name", "UNKNOWN_IMAGE_LOT")).upper().replace("_", " ").strip()
        try:
            budget = int(data.get("recommended_budget_cap", 1500))
        except (ValueError, TypeError):
            budget = 1500
        extracted_specs = str(data.get("extracted_specs", "No detailed specifications could be parsed from the image."))
        
        return {
            "item_name": item_name,
            "recommended_budget_cap": budget,
            "extracted_specs": extracted_specs
        }
    except Exception as e:
        print(f"[FAL VLM FALLBACK] Error during fal VLM extraction: {e}")
        # Robust dictionary fallback structure to prevent server thread runtime errors
        return {
            "item_name": "UNSPECIFIED IMAGE LOT",
            "recommended_budget_cap": 1500,
            "extracted_specs": f"Failed to parse or serialize fal response. Error: {str(e)}"
        }


def ingest_image_rfq_via_fal_bagel(image_url: str) -> dict:
    """
    Connects to the fal-ai/bagel/understand API using the fal_client Python SDK.
    Instructs the model to analyze an image of a physical spec sheet and output ONLY a raw JSON string matching the exact schema.
    """
    import fal_client
    print(f"[BAGEL INTERCEPT] Processing image parameter extraction via fal-ai/bagel/understand...")
    
    prompt = (
        "You are an expert industrial parameter extractor. Analyze the provided image of an industrial component, specification sheet, or document.\n"
        "Extract the parameters and output a clean, un-wrapped JSON object matching this exact schema:\n"
        '{"item_name": "UPPERCASE_ASSET_CODE_WITH_DASHES", "recommended_budget_cap": 1200, "extracted_specs": "Technical parameters summary listing tolerances, dimensions, etc."}\n\n'
        "CRITICAL INSTRUCTIONS:\n"
        "1. Do not use markdown code blocks or wrap your response in triple backticks.\n"
        "2. If the image is unrelated, placeholder, or not a physical specification sheet (e.g. a fantasy castle, scenic photo, or generic image), you MUST STILL OUTPUT A VALID JSON matching the schema using generic placeholder values. For example: "
        '{"item_name": "CASTLE-STRUCTURE-LOT", "recommended_budget_cap": 1500, "extracted_specs": "Scenic fantasy castle structure"}. Do not fail.\n'
        "3. Keep your thinking process extremely short and concise (under 20 words), and make sure the final text ends with the complete, valid, parseable JSON object."
    )
    
    try:
        # Submit the request to fal-ai/bagel/understand
        handler = fal_client.submit(
            "fal-ai/bagel/understand",
            arguments={
                "image_url": image_url,
                "prompt": prompt,
            }
        )
        
        # Block and retrieve the response
        result = handler.get()
        print(f"[BAGEL SUCCESS] Raw response: {result}")
        
        raw_text = ""
        if isinstance(result, dict):
            raw_text = result.get("output", result.get("text", result.get("result", "")))
        else:
            raw_text = str(result)
            
        # Extract and parse the JSON using robust helper
        data = extract_json_from_text(raw_text)
        if data is None:
            raise ValueError("Could not extract JSON from raw Fal Bagel response.")
        item_name = str(data.get("item_name", "UNKNOWN_IMAGE_LOT")).upper().replace("_", "-").strip()
        try:
            budget = int(data.get("recommended_budget_cap", 1500))
        except (ValueError, TypeError):
            budget = 1500
        extracted_specs = str(data.get("extracted_specs", "No detailed specifications could be parsed from the image."))
        
        return {
            "item_name": item_name,
            "recommended_budget_cap": budget,
            "extracted_specs": extracted_specs
        }
    except Exception as e:
        print(f"[BAGEL FALLBACK] Error during fal bagel extraction: {e}")
        # Robust dictionary fallback structure to prevent server thread runtime errors
        return {
            "item_name": "UNSPECIFIED IMAGE LOT",
            "recommended_budget_cap": 1500,
            "extracted_specs": f"Failed to parse or serialize fal bagel response. Error: {str(e)}"
        }


